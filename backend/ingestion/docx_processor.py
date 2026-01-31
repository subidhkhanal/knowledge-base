from pathlib import Path
from typing import List, Dict, Any
import io

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class DOCXProcessor:
    """Process Microsoft Word DOCX files."""

    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError(
                "python-docx is not installed. Install with: pip install python-docx"
            )

    def process(self, file_path: str | Path) -> List[Dict[str, Any]]:
        """
        Extract text from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            List of dicts with 'text' and metadata
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")

        doc = Document(str(file_path))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            return []

        return [{
            "text": full_text.strip(),
            "source": file_path.name,
            "source_type": "docx",
            "page": None
        }]

    def process_bytes(self, content: bytes, filename: str) -> List[Dict[str, Any]]:
        """
        Process DOCX from bytes (for file uploads).

        Args:
            content: DOCX file content as bytes
            filename: Original filename

        Returns:
            List of dicts with 'text' and metadata
        """
        doc = Document(io.BytesIO(content))

        paragraphs = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        paragraphs.append(text)

        full_text = "\n\n".join(paragraphs)

        if not full_text.strip():
            return []

        return [{
            "text": full_text.strip(),
            "source": filename,
            "source_type": "docx",
            "page": None
        }]


def is_docx_available() -> bool:
    """Check if python-docx is available."""
    return DOCX_AVAILABLE
